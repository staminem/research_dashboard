from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from pdf2docx import Converter
from pdf2docx import parse
from io import StringIO
output_string = StringIO()
from nltk import word_tokenize
from pdfminer.layout import LAParams
from docx import Document
from docx.shared import Length
document = Document("/Users/stefangirsberger/Desktop/Papers/20180300_0-00015.docx")
from pdfminer.high_level import extract_pages, extract_text
from pdfminer.layout import LTTextContainer


text = extract_text("/Users/stefangirsberger/Desktop/Papers/20180300_0-00015.pdf")
text_norefs = text.split('References')[0]
text_file = write('text_norefs.txt')

"""headings = []
paragraphs = []
for para in document.paragraphs:
    paragraphs.append(para.text)
for content in document.paragraphs:
    if content.style.name=="Heading 1" or content.style.name=="Heading 2" or content.style.name=="Heading 3":
        headings.append(content.text)"""


parser = PlaintextParser.from_string(text_norefs, Tokenizer("english"))
summariser = LsaSummarizer()
summary = summariser(parser.document, 20)
for sentence in summary:
    print(sentence)
breakpoint()
document = Document()
summary_text = word_tokenize(summary)
print(summary_text)
breakpoint()
document.add_paragraph()
document.save("/Users/stefangirsberger/Desktop/Papers/20180300_0-00015_summarised.docx")

